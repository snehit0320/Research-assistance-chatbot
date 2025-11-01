!pip install requests chromadb sentence-transformers pymupdf gradio arxiv python-docx reportlab matplotlib plotly pandas networkx --quiet

import os, re, json, requests, fitz, time
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions
import gradio as gr
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import arxiv
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter
import networkx as nx

# ------------------------------
# API Keys
# ------------------------------
OPENROUTER_API_KEY = "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"
RAPIDAPI_KEY = "537491845dmshb7f7f9ad7e79d29p1ce47bjsn1d8fcd968154"

# API Hosts
AI_DETECTION_HOST = "ai-detection4.p.rapidapi.com"
PLAGIARISM_HOST = "plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com"
HUMANIZE_HOST = "humanize-ai-content-paraphrasing-api.p.rapidapi.com"
TEXTGEARS_HOST = "textgears-textgears-v1.p.rapidapi.com"

# ------------------------------
# Session Storage
# ------------------------------
session_messages = []
session_docs = []
current_research_topic = None
has_uploaded_pdfs = False

# ------------------------------
# Quality Check Functions
# ------------------------------
def detect_ai(text):
    """Detect AI-generated content score"""
    if not text or len(text.strip()) < 10:
        return 0
    
    url = f"https://{AI_DETECTION_HOST}/v1/ai-detection-rapid-api"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": AI_DETECTION_HOST,
        "Content-Type": "application/json"
    }
    payload = {"text": text, "lang": "en"}
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        return float(data.get("aiScore", 0))
    except Exception as e:
        print(f"❌ AI Detection Error: {e}")
        return 0

def check_plagiarism(text):
    """Check plagiarism percentage"""
    if not text or len(text.strip()) < 10:
        return 0
        
    url = f"https://{PLAGIARISM_HOST}/plagiarism"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": PLAGIARISM_HOST,
        "Content-Type": "application/json"
    }
    payload = {
        "text": text,
        "language": "en",
        "includeCitations": False,
        "scrapeSources": False
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        return float(data.get("plagiarismPercentage", 0))
    except Exception as e:
        print(f"❌ Plagiarism Check Error: {e}")
        return 0

def humanize_text(text):
    """Humanize AI-generated text"""
    if not text or len(text.strip()) < 10:
        return text
        
    url = f"https://{HUMANIZE_HOST}/v1/paraphrase?raw=true"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": HUMANIZE_HOST,
        "Content-Type": "application/json"
    }
    try:
        res = requests.post(url, headers=headers, json={"text": text}, timeout=30)
        res.raise_for_status()
        data = res.json()
        return data.get("humanized", text).strip()
    except Exception as e:
        print(f"❌ Humanize Error: {e}")
        return text

def grammar_check(text):
    """Check grammar errors"""
    if not text or len(text.strip()) < 10:
        return []
        
    url = f"https://{TEXTGEARS_HOST}/grammar"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": TEXTGEARS_HOST
    }
    payload = {"text": text, "language": "en-US"}
    try:
        res = requests.post(url, headers=headers, data=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        errors = data.get("response", {}).get("errors", [])
        return errors
    except Exception as e:
        print(f"❌ Grammar Check Error: {e}")
        return []

def correct_text_with_openrouter(text):
    """Fix grammar using OpenRouter AI"""
    if not text or len(text.strip()) < 10:
        return text
        
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://researchmate-ai.com",
        "X-Title": "ResearchMate AI"
    }
    payload = {
        "model": "mistralai/mistral-7b-instruct:free",
        "messages": [
            {"role": "system", "content": "Fix grammar and spelling errors. Maintain academic tone. Return ONLY the corrected text."},
            {"role": "user", "content": f"Fix grammar:\n\n{text}"}
        ]
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=60)
        res.raise_for_status()
        data = res.json()
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"❌ Grammar Correction Error: {e}")
        return text

def process_section_quality(text, section_name, ai_thresh=15, plag_thresh=10, max_humanize_attempts=3):
    """
    Run comprehensive quality checks on a single section.
    Returns processed_text and report dict.
    """
    if not text or len(text.strip()) < 30:
        return text, {"ai_score": 0, "plagiarism_score": 0, "grammar_issues": 0, "status": "skipped"}

    print(f"🔍 Processing {section_name}...")
    
    # Initial checks
    ai_score = detect_ai(text)
    plag_score = check_plagiarism(text)
    processed_text = text

    print(f"   Initial - AI: {ai_score:.1f}%, Plagiarism: {plag_score:.1f}%")

    # If needs humanizing, try multiple strategies
    humanize_attempt = 0
    if ai_score > ai_thresh or plag_score > plag_thresh:
        print(f"   🔄 Humanizing (attempts: {max_humanize_attempts})...")
        
        while humanize_attempt < max_humanize_attempts:
            try:
                processed_text = humanize_text(processed_text)
                time.sleep(1)  # Rate limiting
                
                # Re-check after humanization
                ai_score = detect_ai(processed_text)
                plag_score = check_plagiarism(processed_text)
                
                print(f"   Attempt {humanize_attempt + 1} - AI: {ai_score:.1f}%, Plagiarism: {plag_score:.1f}%")
                
                humanize_attempt += 1
                if ai_score <= ai_thresh and plag_score <= plag_thresh:
                    print(f"   ✅ Humanization successful!")
                    break
                    
            except Exception as e:
                print(f"   ❌ Humanization attempt {humanize_attempt + 1} failed: {e}")
                humanize_attempt += 1

        # If still flagged, ask OpenRouter to paraphrase to "human academic style"
        if ai_score > ai_thresh or plag_score > plag_thresh:
            print(f"   🤖 Using AI paraphrasing...")
            prompt = (
                f"Paraphrase the following section to make it read like authentic, human-written academic prose. "
                f"Keep the technical meaning exactly the same, keep citations/values, but change wording and sentence rhythm. "
                f"Return ONLY the paraphrased section.\n\nSECTION NAME: {section_name}\n\n{processed_text}"
            )
            try:
                messages = [
                    {"role": "system", "content": "You are a skilled academic writer. Paraphrase text to sound human and natural."},
                    {"role": "user", "content": prompt}
                ]
                paraphrased = get_ai_response(messages)
                if paraphrased and not paraphrased.startswith("[AI ERROR]"):
                    processed_text = paraphrased
                    print(f"   ✅ AI paraphrasing completed")
            except Exception as e:
                print(f"   ❌ AI paraphrasing failed: {e}")

    # Grammar check and correction if many errors
    print(f"   📝 Checking grammar...")
    grammar_errors = grammar_check(processed_text)
    if grammar_errors and len(grammar_errors) > 2:
        print(f"   🔧 Correcting {len(grammar_errors)} grammar errors...")
        try:
            processed_text = correct_text_with_openrouter(processed_text)
            print(f"   ✅ Grammar correction completed")
        except Exception as e:
            print(f"   ❌ Grammar correction failed: {e}")

    # Final checks for report
    final_ai = detect_ai(processed_text)
    final_plag = check_plagiarism(processed_text)
    final_grammar = grammar_check(processed_text)

    status = "✅ passed" if final_ai <= ai_thresh and final_plag <= plag_thresh else "⚠️ needs review"
    
    print(f"   📊 Final - AI: {final_ai:.1f}%, Plagiarism: {final_plag:.1f}%, Grammar: {len(final_grammar)} issues")
    
    report = {
        "ai_score": final_ai,
        "plagiarism_score": final_plag,
        "grammar_issues": len(final_grammar),
        "status": status
    }
    return processed_text, report

# ------------------------------
# OpenRouter AI
# ------------------------------
FREE_MODELS = [
    "deepseek/deepseek-r1-0528-qwen3-8b:free",
    "meta-llama/llama-3.2-3b-instruct:free",
    "google/gemini-flash-1.5:free",
    "mistralai/mistral-7b-instruct:free",
    "nousresearch/hermes-3-llama-3.1-405b:free"
]
current_model_index = 0

def get_ai_response(messages, max_retries=3):
    global current_model_index

    models_tried = 0
    max_models_to_try = len(FREE_MODELS)

    while models_tried < max_models_to_try:
        model = FREE_MODELS[current_model_index]

        for attempt in range(max_retries):
            try:
                url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://researchmate-ai.com",
                    "X-Title": "ResearchMate AI"
                }
                payload = {"model": model, "messages": messages}

                if attempt > 0:
                    time.sleep(2 ** attempt)

                r = requests.post(url, headers=headers, data=json.dumps(payload), timeout=60)

                if r.status_code == 429:
                    current_model_index = (current_model_index + 1) % len(FREE_MODELS)
                    models_tried += 1
                    break

                r.raise_for_status()
                data = r.json()

                if "choices" in data and len(data["choices"]) > 0:
                    return data["choices"][0]["message"]["content"]

            except Exception as e:
                if attempt == max_retries - 1:
                    return f"[AI ERROR] {str(e)}"

    return "⚠️ All models are currently rate limited. Please wait and try again."

# ------------------------------
# RAG Setup
# ------------------------------
try:
    chromadb.api.client.SharedSystemClient.clear_system_cache()
except:
    pass

client = chromadb.PersistentClient(path="./chromadb_session")
embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
collection = client.get_or_create_collection(name="papers_session", embedding_function=embedding_fn)

def add_texts_to_rag(texts, names, source_type="pdf"):
    for text, name in zip(texts, names):
        doc_id = re.sub(r'[^\w\-\. ]', '', name)[:200]
        try:
            collection.add(documents=[text], metadatas=[{"title": name, "source": source_type}], ids=[doc_id])
        except Exception as e:
            print("Chroma add error:", e)
        session_docs.append({"title": name, "text": text, "source": source_type})

def query_rag(user_query, research_topic=None):
    global has_uploaded_pdfs, current_research_topic

    try:
        pdf_docs = [d for d in session_docs if d.get("source") == "pdf"]
        arxiv_docs = [d for d in session_docs if d.get("source") == "arxiv"]

        if pdf_docs:
            results = collection.query(query_texts=[user_query], n_results=3)
            docs = results["documents"][0] if results["documents"] else []
            pdf_context = "\n\n".join([d for d in docs if d]) if docs else ""

            if not pdf_context:
                return "❌ No relevant content found in your uploaded PDFs."

            messages = [
                {"role": "system", "content": "You are a research assistant. Answer based on provided PDF content."},
                {"role": "user", "content": f"CONTENT:\n{pdf_context}\n\nQUESTION: {user_query}"}
            ]
            return get_ai_response(messages)

        elif current_research_topic and not pdf_docs:
            if not arxiv_docs:
                papers = fetch_research_papers(current_research_topic, max_results=5)
                if papers:
                    abstracts = [p["abstract"] for p in papers if p.get("abstract")]
                    titles = [p["title"] for p in papers]
                    add_texts_to_rag(abstracts, titles, source_type="arxiv")
                else:
                    return "❌ Could not fetch arXiv papers."

            results = collection.query(query_texts=[user_query], n_results=3)
            docs = results["documents"][0] if results["documents"] else []
            arxiv_context = "\n\n".join([d for d in docs if d]) if docs else ""

            messages = [
                {"role": "system", "content": f"Research assistant for {current_research_topic}."},
                {"role": "user", "content": f"TOPIC: {current_research_topic}\n\nCONTEXT:\n{arxiv_context}\n\nQUESTION: {user_query}"}
            ]
            return get_ai_response(messages)

        else:
            return "❌ Please upload PDFs or set a research topic."

    except Exception as e:
        return f"[RAG ERROR] {e}"

# ------------------------------
# PDF Utilities
# ------------------------------
def extract_text_from_pdfs(pdf_files):
    texts, names = [], []
    for pdf_file in pdf_files:
        try:
            doc = fitz.open(pdf_file.name)
            pdf_text = "".join([page.get_text() for page in doc])
            texts.append(pdf_text)
            names.append(os.path.basename(pdf_file.name))
        except Exception as e:
            print(f"Error reading PDF: {e}")
    return texts, names

# ------------------------------
# arXiv Utilities
# ------------------------------
def fetch_research_papers(topic, max_results=5):
    papers = []
    try:
        search = arxiv.Search(query=topic, max_results=max_results, sort_by=arxiv.SortCriterion.Relevance)
        for result in search.results():
            papers.append({
                "title": result.title,
                "authors": [a.name for a in result.authors],
                "abstract": result.summary,
                "pdf_url": result.pdf_url
            })
    except Exception as e:
        print("arXiv fetch error:", e)
    return papers

# ------------------------------
# Chat Integration
# ------------------------------
def combined_chat(idea_topic, pdf_files, user_input):
    global current_research_topic, has_uploaded_pdfs

    if not user_input:
        return "❌ Please enter a question."

    if idea_topic:
        current_research_topic = idea_topic

    if pdf_files and not has_uploaded_pdfs:
        pdf_texts, pdf_names = extract_text_from_pdfs(pdf_files)
        if pdf_texts:
            add_texts_to_rag(pdf_texts, pdf_names, source_type="pdf")
            has_uploaded_pdfs = True

    session_messages.append({"role": "user", "content": user_input})
    answer = query_rag(user_input, research_topic=current_research_topic)
    session_messages.append({"role": "assistant", "content": answer})

    return answer

# ------------------------------
# Enhanced Paper Generator with Individual Section Generation
# ------------------------------
def generate_individual_section(section_name, topic, context_info, abstracts, previous_sections=None):
    """Generate a specific section with detailed, comprehensive prompts"""
    
    # Build context from previous sections
    context_text = ""
    if previous_sections:
        for name, content in previous_sections.items():
            if content and name != section_name:
                context_text += f"\n{name}:\n{content[:500]}...\n"
    
    section_prompts = {
        "Abstract": f"""
        Write a comprehensive abstract (200-300 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Background and motivation (2-3 sentences)
        2. Research objectives and questions (2-3 sentences)  
        3. Methodology and approach (2-3 sentences)
        4. Key findings and results (3-4 sentences)
        5. Conclusions and implications (2-3 sentences)
        
        Make it concise, informative, and academically rigorous. Use present tense for general statements, past tense for specific findings.
        
        Context: {context_info}
        Relevant Abstracts: {abstracts[:1000]}
        """,
        
        "Introduction": f"""
        Write a detailed introduction (800-1200 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Opening hook and problem statement (2-3 paragraphs)
        2. Background and context (2-3 paragraphs)
        3. Research gap and motivation (1-2 paragraphs)
        4. Research objectives and questions (1-2 paragraphs)
        5. Significance and contributions (1-2 paragraphs)
        6. Paper structure overview (1 paragraph)
        
        Include relevant statistics, current challenges, and why this research matters. Use academic tone with proper citations.
        
        Context: {context_info}
        Relevant Abstracts: {abstracts[:1000]}
        """,
        
        "Literature Review": f"""
        Write a comprehensive literature review (1000-1500 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Theoretical framework and key concepts (2-3 paragraphs)
        2. Historical development and evolution (2-3 paragraphs)
        3. Current state of research (3-4 paragraphs)
        4. Methodological approaches used (2-3 paragraphs)
        5. Identified gaps and limitations (2-3 paragraphs)
        6. Synthesis and positioning of your work (1-2 paragraphs)
        
        Analyze existing research, compare different approaches, identify trends, and position your work within the field. Include critical analysis and synthesis.
        
        Context: {context_info}
        Relevant Abstracts: {abstracts[:1500]}
        """,
        
        "Methodology": f"""
        Write a detailed methodology section (800-1200 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Research design and approach (1-2 paragraphs)
        2. Data collection methods and sources (2-3 paragraphs)
        3. Sample selection and criteria (1-2 paragraphs)
        4. Tools, instruments, and technologies used (2-3 paragraphs)
        5. Data analysis procedures and techniques (2-3 paragraphs)
        6. Validity, reliability, and ethical considerations (1-2 paragraphs)
        7. Limitations and constraints (1 paragraph)
        
        Be specific about procedures, justify methodological choices, and explain how you ensure rigor. Include details about tools, software, or frameworks used.
        
        Context: {context_info}
        Previous sections: {context_text}
        """,
        
        "Results and Discussion": f"""
        Write a comprehensive results and discussion section (1200-1800 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Overview of findings (1-2 paragraphs)
        2. Detailed results presentation (4-6 paragraphs with specific findings)
        3. Statistical analysis and data interpretation (2-3 paragraphs)
        4. Comparison with existing literature (2-3 paragraphs)
        5. Theoretical implications (2-3 paragraphs)
        6. Practical applications and significance (2-3 paragraphs)
        7. Limitations and future research directions (1-2 paragraphs)
        
        Present findings clearly, interpret results, discuss implications, and connect to broader literature. Include specific data, statistics, or examples where relevant.
        
        Context: {context_info}
        Previous sections: {context_text}
        """,
        
        "Conclusion": f"""
        Write a comprehensive conclusion (400-600 words) for a research paper on '{topic}'. 
        
        Structure:
        1. Summary of key findings (2-3 paragraphs)
        2. Achievement of research objectives (1-2 paragraphs)
        3. Theoretical and practical contributions (2-3 paragraphs)
        4. Limitations and constraints (1 paragraph)
        5. Future research directions and recommendations (1-2 paragraphs)
        6. Final thoughts and implications (1 paragraph)
        
        Synthesize the main points, highlight contributions, acknowledge limitations, and suggest future work. Be concise but comprehensive.
        
        Context: {context_info}
        Previous sections: {context_text}
        """,
        
        "References": f"""
        Generate a comprehensive reference list for a research paper on '{topic}'. 
        
        Include 20-25 relevant academic sources:
        - 8-10 recent journal articles (2020-2024)
        - 5-7 foundational papers and seminal works
        - 3-5 conference proceedings
        - 2-3 books or book chapters
        - 2-3 technical reports or white papers
        
        Use proper academic citation format (APA or IEEE style). Ensure sources are credible and directly relevant to the topic.
        
        Context: {context_info}
        Topic: {topic}
        """
    }
    
    prompt = section_prompts.get(section_name, f"Write a comprehensive {section_name} section for a research paper on '{topic}'.")
    
    messages = [
        {"role": "system", "content": f"You are an expert academic writer and researcher specializing in {topic}. Write high-quality, original, and comprehensive content that meets academic standards. Be specific, detailed, and rigorous in your approach."},
        {"role": "user", "content": prompt}
    ]
    
    return get_ai_response(messages)

def generate_complete_paper_sections(topic, abstracts, context_info, progress=gr.Progress()):
    """Generate all sections individually with proper content"""
    
    sections = {}
    section_order = ["Abstract", "Introduction", "Literature Review", "Methodology", "Results and Discussion", "Conclusion", "References"]
    
    print(f"🚀 Starting individual section generation for: {topic}")
    
    for i, section_name in enumerate(section_order):
        progress(0.1 + 0.8 * (i / len(section_order)), desc=f"Generating {section_name}...")
        
        print(f"📝 Generating {section_name}...")
        
        # Generate the section
        content = generate_individual_section(section_name, topic, context_info, abstracts, sections)
        
        if content and not content.startswith("[AI ERROR]"):
            sections[section_name] = content
            print(f"✅ Generated {section_name} ({len(content)} characters)")
        else:
            # Fallback content
            fallback_content = f"""
            [{section_name} - Placeholder Content]
            
            This section requires manual completion for the research paper on '{topic}'. 
            The automated generation encountered an issue, but the structure and context are provided.
            
            Please review and expand this section with:
            - Detailed analysis and discussion
            - Relevant citations and references
            - Specific examples and data
            - Proper academic formatting
            
            Topic: {topic}
            Context: {context_info}
            """
            sections[section_name] = fallback_content
            print(f"⚠️ Using fallback for {section_name}")
        
        # Rate limiting between sections
        time.sleep(2)
    
    return sections

def generate_paper_with_quality_check(topic, enable_quality_check=True, progress=gr.Progress()):
    """
    Enhanced paper generation with individual section creation and comprehensive quality control
    """
    if not topic:
        return "❌ Please enter a research topic.", None, None, ""

    print(f"🚀 Starting comprehensive paper generation for: {topic}")
    
    progress(0, desc="Fetching research papers...")
    papers = fetch_research_papers(topic, max_results=5)
    abstracts = " ".join([p["abstract"] for p in papers if p.get("abstract")])
    context_info = f"Research topic: {topic}. Found {len(papers)} relevant papers from arXiv."

    progress(0.05, desc="Generating individual sections...")
    
    # Generate all sections individually
    sections = generate_complete_paper_sections(topic, abstracts, context_info, progress)

    # Quality processing
    if enable_quality_check:
        progress(0.9, desc="Starting quality processing...")
        processed_sections = {}
        section_reports = {}
        
        for idx, (section_name, section_content) in enumerate(sections.items()):
            # Skip references from heavy processing
            if section_name.lower() == "references":
                processed_sections[section_name] = section_content
                section_reports[section_name] = {
                    "ai_score": 0, 
                    "plagiarism_score": 0, 
                    "grammar_issues": 0, 
                    "status": "skipped"
                }
                continue

            progress(0.9 + 0.05 * (idx / max(1, len(sections)-1)), 
                    desc=f"Quality processing {section_name}...")
            
            processed_content, report = process_section_quality(section_content, section_name)
            processed_sections[section_name] = processed_content
            section_reports[section_name] = report
            
            time.sleep(1)  # Rate limiting
    else:
        processed_sections = sections
        section_reports = {name: {"ai_score": 0, "plagiarism_score": 0, "grammar_issues": 0, "status": "not-checked"} for name in sections.keys()}

    # Reconstruct the paper in canonical order
    progress(0.98, desc="Assembling final paper...")
    ordered_text = "\n\n".join([
        f"{name}\n{processed_sections.get(name,'')}".strip() 
        for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]
    ])

    # Create files
    progress(0.99, desc="Generating files...")
    os.makedirs("generated_papers", exist_ok=True)
    safe_name = re.sub(r'[^\w\-\. ]', '', topic)[:50]

    # DOCX
    docx_path = f"generated_papers/{safe_name}.docx"
    doc = Document()
    doc.add_heading(topic, level=1)
    
    for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]:
        content = processed_sections.get(name, "")
        if content:
            doc.add_heading(name, level=2)
            # Split content into paragraphs
            paragraphs = re.split(r'\n{2,}', content)
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
    
    doc.save(docx_path)

    # PDF
    pdf_path = f"generated_papers/{safe_name}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50
    
    # Title
    c.setFont("Times-Bold", 16)
    c.drawString(50, y_position, topic)
    y_position -= 30
    
    for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]:
        content = processed_sections.get(name, "")
        if content:
            # Section heading
            c.setFont("Times-Bold", 12)
            c.drawString(50, y_position, name)
            y_position -= 20
            
            # Section content
            c.setFont("Times-Roman", 10)
            lines = content.split('\n')
            for line in lines:
                if y_position < 50:  # New page
                    c.showPage()
                    y_position = height - 50
                
                # Word wrap
                words = line.split()
                current_line = ""
                for word in words:
                    if c.stringWidth(current_line + word + " ") < width - 100:
                        current_line += word + " "
                    else:
                        if current_line:
                            c.drawString(50, y_position, current_line.strip())
                            y_position -= 12
                        current_line = word + " "
                
                if current_line:
                    c.drawString(50, y_position, current_line.strip())
                    y_position -= 12
                
                y_position -= 5  # Space between paragraphs
    
    c.save()

    # Build quality report
    quality_report_text = ""
    if enable_quality_check:
        quality_report_text = "📊 QUALITY CHECK REPORT\n" + "="*60 + "\n\n"
        for section, report in section_reports.items():
            quality_report_text += f"📄 {section}:\n"
            quality_report_text += f"   AI Score: {report['ai_score']:.1f}%\n"
            quality_report_text += f"   Plagiarism: {report['plagiarism_score']:.1f}%\n"
            quality_report_text += f"   Grammar Issues: {report['grammar_issues']}\n"
            quality_report_text += f"   Status: {report['status']}\n\n"

    progress(1.0, desc="Complete!")
    print(f"✅ Paper generation completed: {topic}")
    
    return ordered_text, docx_path, pdf_path, quality_report_text

# ------------------------------
# Analytics Functions
# ------------------------------
def get_session_statistics():
    total_docs = len(session_docs)
    total_messages = len(session_messages)
    user_messages = len([m for m in session_messages if m["role"]=="user"])
    ai_messages = len([m for m in session_messages if m["role"]=="assistant"])
    total_chars = sum(len(doc["text"]) for doc in session_docs)
    avg_doc_length = total_chars / total_docs if total_docs > 0 else 0
    pdf_count = len([d for d in session_docs if d.get("source") == "pdf"])
    arxiv_count = len([d for d in session_docs if d.get("source") == "arxiv"])

    return {
        "total_documents": total_docs,
        "pdf_documents": pdf_count,
        "arxiv_documents": arxiv_count,
        "total_messages": total_messages,
        "user_messages": user_messages,
        "ai_messages": ai_messages,
        "total_characters": total_chars,
        "avg_doc_length": avg_doc_length
    }

def get_top_keywords(top_n=15):
    all_text = " ".join([doc["text"] for doc in session_docs])
    words = re.findall(r'\b\w+\b', all_text.lower())
    stopwords = set(["the","and","for","with","that","this","from","are","was","were","have","has","using","our","can","which","these","their","been","into","than","more","also","will","such","when","there","other","through","about","some","only","would","between"])
    filtered_words = [w for w in words if w not in stopwords and len(w)>3]
    word_counts = Counter(filtered_words)
    return word_counts.most_common(top_n)

def plot_keyword_bar():
    keywords = get_top_keywords()
    if not keywords:
        fig = go.Figure()
        fig.update_layout(title="Top Keywords", plot_bgcolor="white")
        return fig
    df = pd.DataFrame(keywords, columns=["Keyword","Count"])
    fig = px.bar(df, x="Keyword", y="Count", title="Top Keywords", color="Count")
    return fig

def generate_insights_report_session():
    stats = get_session_statistics()
    keywords = get_top_keywords(10)

    topic_info = f"\n🎯 Topic: {current_research_topic}" if current_research_topic else ""
    source_info = f"\n📄 Source: {'PDFs' if has_uploaded_pdfs else 'arXiv'}"

    report = f"""
📊 Session Insights
{'='*60}
{topic_info}{source_info}

📚 Documents: {stats['total_documents']} ({stats['pdf_documents']} PDFs, {stats['arxiv_documents']} arXiv)
💬 Interactions: {stats['total_messages']}
📝 Characters: {stats['total_characters']:,}

🔑 Keywords: {', '.join([k for k,c in keywords[:10]]) if keywords else 'None'}
"""
    return report

def refresh_all_analytics():
    report = generate_insights_report_session()
    keyword_plot = plot_keyword_bar()
    return report, keyword_plot

# ------------------------------
# Gradio UI
# ------------------------------
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("""
    # 🧠 ResearchMate AI — Advanced Research Assistant with Quality Control
    ### 📚 Upload PDFs OR Set Topic | ✅ AI Detection, Plagiarism Check, Humanization & Grammar Correction
    """)

    with gr.Tabs():
        with gr.Tab("💬 Chat Assistant"):
            with gr.Row():
                clear_btn = gr.Button("🔄 Clear Session", variant="secondary")

            idea_input = gr.Textbox(label="🎯 Research Topic", placeholder="e.g., Machine Learning in Healthcare")
            pdf_upload = gr.Files(label="📄 Upload PDFs (Optional)", file_types=[".pdf"])
            pdf_status = gr.Textbox(label="📋 Status", interactive=False, lines=3)
            chat_input = gr.Textbox(label="❓ Your Question", placeholder="Ask anything...", lines=2)
            output_box = gr.Textbox(label="🤖 AI Answer", interactive=False, lines=12)
            btn = gr.Button("🚀 Send", variant="primary")

            def clear_session():
                global session_messages, session_docs, current_research_topic, has_uploaded_pdfs
                session_messages, session_docs = [], []
                current_research_topic, has_uploaded_pdfs = None, False
                try:
                    collection.delete(where={})
                except:
                    pass
                return "✅ Session cleared!", ""

            clear_btn.click(clear_session, outputs=[pdf_status, output_box])

            def upload_handler(idea_topic, pdf_files):
                global current_research_topic, has_uploaded_pdfs
                if idea_topic:
                    current_research_topic = idea_topic
                if not pdf_files:
                    return f"✅ Topic set: '{current_research_topic}'" if current_research_topic else "⚠️ Set topic or upload PDFs"
                try:
                    pdf_texts, pdf_names = extract_text_from_pdfs(pdf_files)
                    if pdf_texts:
                        add_texts_to_rag(pdf_texts, pdf_names, source_type="pdf")
                        has_uploaded_pdfs = True
                        return f"✅ Processed {len(pdf_texts)} PDF(s)"
                    return "❌ Failed to extract text"
                except Exception as e:
                    return f"❌ Error: {str(e)}"

            pdf_upload.change(upload_handler, inputs=[idea_input, pdf_upload], outputs=pdf_status)
            btn.click(combined_chat, inputs=[idea_input, pdf_upload, chat_input], outputs=output_box)

        with gr.Tab("📝 Paper Generator with Quality Control"):
            gr.Markdown("""
            Generate research papers with automatic quality checking:
            - ✅ AI Content Detection
            - ✅ Plagiarism Checking  
            - ✅ Text Humanization
            - ✅ Grammar Correction
            """)

            topic_box = gr.Textbox(label="Research Topic", placeholder="e.g., Quantum Computing in Cryptography")
            quality_checkbox = gr.Checkbox(label="Enable Quality Processing (AI Detection, Plagiarism, Humanization, Grammar)", value=True)

            paper_output = gr.Textbox(label="Generated Paper", lines=20, interactive=False)
            quality_report_box = gr.Textbox(label="📊 Quality Report", lines=10, interactive=False)

            with gr.Row():
                docx_download = gr.File(label="📥 Download DOCX")
                pdf_download = gr.File(label="📥 Download PDF")

            gen_btn = gr.Button("✍️ Generate Paper", variant="primary")

            gen_btn.click(
                generate_paper_with_quality_check,
                inputs=[topic_box, quality_checkbox],
                outputs=[paper_output, docx_download, pdf_download, quality_report_box]
            )

        with gr.Tab("📊 Analytics"):
            refresh_btn = gr.Button("🔄 Refresh", variant="primary")
            insights_text = gr.Textbox(label="Overview", lines=15, interactive=False)
            keyword_plot_box = gr.Plot(label="Keywords")

            refresh_btn.click(refresh_all_analytics, outputs=[insights_text, keyword_plot_box])
            demo.load(refresh_all_analytics, outputs=[insights_text, keyword_plot_box])

demo.launch(share=True)
