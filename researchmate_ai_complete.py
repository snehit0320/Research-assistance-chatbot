import os, re, json, requests, fitz, time
from dotenv import load_dotenv
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions
import gradio as gr
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import arxiv
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter
import networkx as nx

# Load environment variables from .env file
load_dotenv()

# ------------------------------
# API Keys (loaded from environment variables)
# ------------------------------
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
RAPIDAPI_KEY = os.getenv("RAPIDAPI_KEY", "")

# Validate API keys are set
if not OPENROUTER_API_KEY or not RAPIDAPI_KEY:
    print("⚠️ WARNING: API keys not found in environment variables!")
    print("Please create a .env file with your API keys. See env.example for reference.")
    raise ValueError("API keys must be set in environment variables. Create a .env file based on env.example")

# API Hosts (can be overridden via environment variables)
AI_DETECTION_HOST = os.getenv("AI_DETECTION_HOST", "ai-detection4.p.rapidapi.com")
PLAGIARISM_HOST = os.getenv("PLAGIARISM_HOST", "plagiarism-checker-and-auto-citation-generator-multi-lingual.p.rapidapi.com")
HUMANIZE_HOST = os.getenv("HUMANIZE_HOST", "humanize-ai-content-paraphrasing-api.p.rapidapi.com")
TEXTGEARS_HOST = os.getenv("TEXTGEARS_HOST", "textgears-textgears-v1.p.rapidapi.com")

# ------------------------------
# Session Storage
# ------------------------------
session_messages = []
session_docs = []
current_research_topic = None
has_uploaded_pdfs = False

# ------------------------------
# Quality Check Functions
# ------------------------------
def detect_ai(text):
    """Detect AI-generated content score"""
    if not text or len(text.strip()) < 10:
        return 0
    
    url = f"https://{AI_DETECTION_HOST}/v1/ai-detection-rapid-api"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": AI_DETECTION_HOST,
        "Content-Type": "application/json"
    }
    payload = {"text": text, "lang": "en"}
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        return float(data.get("aiScore", 0))
    except Exception as e:
        print(f"❌ AI Detection Error: {e}")
        return 0

def check_plagiarism(text):
    """Check plagiarism percentage"""
    if not text or len(text.strip()) < 10:
        return 0
        
    url = f"https://{PLAGIARISM_HOST}/plagiarism"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": PLAGIARISM_HOST,
        "Content-Type": "application/json"
    }
    payload = {
        "text": text,
        "language": "en",
        "includeCitations": False,
        "scrapeSources": False
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        return float(data.get("plagiarismPercentage", 0))
    except Exception as e:
        print(f"❌ Plagiarism Check Error: {e}")
        return 0

def humanize_text(text):
    """Humanize AI-generated text"""
    if not text or len(text.strip()) < 10:
        return text
        
    url = f"https://{HUMANIZE_HOST}/v1/paraphrase?raw=true"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": HUMANIZE_HOST,
        "Content-Type": "application/json"
    }
    try:
        res = requests.post(url, headers=headers, json={"text": text}, timeout=30)
        res.raise_for_status()
        data = res.json()
        return data.get("humanized", text).strip()
    except Exception as e:
        print(f"❌ Humanize Error: {e}")
        return text

def grammar_check(text):
    """Check grammar errors"""
    if not text or len(text.strip()) < 10:
        return []
        
    url = f"https://{TEXTGEARS_HOST}/grammar"
    headers = {
        "x-rapidapi-key": RAPIDAPI_KEY,
        "x-rapidapi-host": TEXTGEARS_HOST
    }
    payload = {"text": text, "language": "en-US"}
    try:
        res = requests.post(url, headers=headers, data=payload, timeout=30)
        res.raise_for_status()
        data = res.json()
        errors = data.get("response", {}).get("errors", [])
        return errors
    except Exception as e:
        print(f"❌ Grammar Check Error: {e}")
        return []

def correct_text_with_openrouter(text):
    """Fix grammar using OpenRouter AI"""
    if not text or len(text.strip()) < 10:
        return text
        
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://researchmate-ai.com",
        "X-Title": "ResearchMate AI"
    }
    payload = {
        "model": "mistralai/mistral-7b-instruct:free",
        "messages": [
            {"role": "system", "content": "Fix grammar and spelling errors. Maintain academic tone. Return ONLY the corrected text."},
            {"role": "user", "content": f"Fix grammar:\n\n{text}"}
        ]
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=60)
        res.raise_for_status()
        data = res.json()
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"❌ Grammar Correction Error: {e}")
        return text

def process_section_quality(text, section_name, ai_thresh=15, plag_thresh=10, max_humanize_attempts=3):
    """
    Run comprehensive quality checks on a single section.
    Returns processed_text and report dict.
    """
    if not text or len(text.strip()) < 30:
        return text, {"ai_score": 0, "plagiarism_score": 0, "grammar_issues": 0, "status": "skipped"}

    print(f"🔍 Processing {section_name}...")
    
    # Initial checks
    ai_score = detect_ai(text)
    plag_score = check_plagiarism(text)
    processed_text = text

    print(f"   Initial - AI: {ai_score:.1f}%, Plagiarism: {plag_score:.1f}%")

    # If needs humanizing, try multiple strategies
    humanize_attempt = 0
    if ai_score > ai_thresh or plag_score > plag_thresh:
        print(f"   🔄 Humanizing (attempts: {max_humanize_attempts})...")
        
        while humanize_attempt < max_humanize_attempts:
            try:
                processed_text = humanize_text(processed_text)
                time.sleep(1)  # Rate limiting
                
                # Re-check after humanization
                ai_score = detect_ai(processed_text)
                plag_score = check_plagiarism(processed_text)
                
                print(f"   Attempt {humanize_attempt + 1} - AI: {ai_score:.1f}%, Plagiarism: {plag_score:.1f}%")
                
                humanize_attempt += 1
                if ai_score <= ai_thresh and plag_score <= plag_thresh:
                    print(f"   ✅ Humanization successful!")
                    break
                    
            except Exception as e:
                print(f"   ❌ Humanization attempt {humanize_attempt + 1} failed: {e}")
                humanize_attempt += 1

        # If still flagged, ask OpenRouter to paraphrase to "human academic style"
        if ai_score > ai_thresh or plag_score > plag_thresh:
            print(f"   🤖 Using AI paraphrasing...")
            prompt = (
                f"Paraphrase the following section to make it read like authentic, human-written academic prose. "
                f"Keep the technical meaning exactly the same, keep citations/values, but change wording and sentence rhythm. "
                f"Return ONLY the paraphrased section.\n\nSECTION NAME: {section_name}\n\n{processed_text}"
            )
            try:
                messages = [
                    {"role": "system", "content": "You are a skilled academic writer. Paraphrase text to sound human and natural."},
                    {"role": "user", "content": prompt}
                ]
                paraphrased = get_ai_response(messages)
                if paraphrased and not paraphrased.startswith("[AI ERROR]"):
                    processed_text = paraphrased
                    print(f"   ✅ AI paraphrasing completed")
            except Exception as e:
                print(f"   ❌ AI paraphrasing failed: {e}")

    # Grammar check and correction if many errors
    print(f"   📝 Checking grammar...")
    grammar_errors = grammar_check(processed_text)
    if grammar_errors and len(grammar_errors) > 2:
        print(f"   🔧 Correcting {len(grammar_errors)} grammar errors...")
        try:
            processed_text = correct_text_with_openrouter(processed_text)
            print(f"   ✅ Grammar correction completed")
        except Exception as e:
            print(f"   ❌ Grammar correction failed: {e}")

    # Final checks for report
    final_ai = detect_ai(processed_text)
    final_plag = check_plagiarism(processed_text)
    final_grammar = grammar_check(processed_text)

    status = "✅ passed" if final_ai <= ai_thresh and final_plag <= plag_thresh else "⚠️ needs review"
    
    print(f"   📊 Final - AI: {final_ai:.1f}%, Plagiarism: {final_plag:.1f}%, Grammar: {len(final_grammar)} issues")
    
    report = {
        "ai_score": final_ai,
        "plagiarism_score": final_plag,
        "grammar_issues": len(final_grammar),
        "status": status
    }
    return processed_text, report

# ------------------------------
# OpenRouter AI
# ------------------------------
FREE_MODELS = [
    "deepseek/deepseek-r1-0528-qwen3-8b:free",
    "meta-llama/llama-3.2-3b-instruct:free",
    "google/gemini-flash-1.5:free",
    "mistralai/mistral-7b-instruct:free",
    "nousresearch/hermes-3-llama-3.1-405b:free"
]
current_model_index = 0

def get_ai_response(messages, max_retries=3):
    global current_model_index

    models_tried = 0
    max_models_to_try = len(FREE_MODELS)

    while models_tried < max_models_to_try:
        model = FREE_MODELS[current_model_index]

        for attempt in range(max_retries):
            try:
                url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://researchmate-ai.com",
                    "X-Title": "ResearchMate AI"
                }
                payload = {"model": model, "messages": messages}

                if attempt > 0:
                    time.sleep(2 ** attempt)

                r = requests.post(url, headers=headers, data=json.dumps(payload), timeout=60)

                if r.status_code == 429:
                    current_model_index = (current_model_index + 1) % len(FREE_MODELS)
                    models_tried += 1
                    break

                r.raise_for_status()
                data = r.json()

                if "choices" in data and len(data["choices"]) > 0:
                    return data["choices"][0]["message"]["content"]

            except Exception as e:
                if attempt == max_retries - 1:
                    return f"[AI ERROR] {str(e)}"

    return "⚠️ All models are currently rate limited. Please wait and try again."

# ------------------------------
# RAG Setup
# ------------------------------
try:
    chromadb.api.client.SharedSystemClient.clear_system_cache()
except:
    pass

client = chromadb.PersistentClient(path="./chromadb_session")
embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
collection = client.get_or_create_collection(name="papers_session", embedding_function=embedding_fn)

def add_texts_to_rag(texts, names, source_type="pdf"):
    for text, name in zip(texts, names):
        doc_id = re.sub(r'[^\w\-\. ]', '', name)[:200]
        try:
            collection.add(documents=[text], metadatas=[{"title": name, "source": source_type}], ids=[doc_id])
        except Exception as e:
            print("Chroma add error:", e)
        session_docs.append({"title": name, "text": text, "source": source_type})

def query_rag(user_query, research_topic=None):
    global has_uploaded_pdfs, current_research_topic

    try:
        pdf_docs = [d for d in session_docs if d.get("source") == "pdf"]
        arxiv_docs = [d for d in session_docs if d.get("source") == "arxiv"]

        if pdf_docs:
            results = collection.query(query_texts=[user_query], n_results=3)
            docs = results["documents"][0] if results["documents"] else []
            pdf_context = "\n\n".join([d for d in docs if d]) if docs else ""

            if not pdf_context:
                return "❌ No relevant content found in your uploaded PDFs."

            messages = [
                {"role": "system", "content": "You are a research assistant. Answer based on provided PDF content."},
                {"role": "user", "content": f"CONTENT:\n{pdf_context}\n\nQUESTION: {user_query}"}
            ]
            return get_ai_response(messages)

        elif current_research_topic and not pdf_docs:
            if not arxiv_docs:
                papers = fetch_research_papers(current_research_topic, max_results=5)
                if papers:
                    abstracts = [p["abstract"] for p in papers if p.get("abstract")]
                    titles = [p["title"] for p in papers]
                    add_texts_to_rag(abstracts, titles, source_type="arxiv")
                else:
                    return "❌ Could not fetch arXiv papers."

            results = collection.query(query_texts=[user_query], n_results=3)
            docs = results["documents"][0] if results["documents"] else []
            arxiv_context = "\n\n".join([d for d in docs if d]) if docs else ""

            messages = [
                {"role": "system", "content": f"Research assistant for {current_research_topic}."},
                {"role": "user", "content": f"TOPIC: {current_research_topic}\n\nCONTEXT:\n{arxiv_context}\n\nQUESTION: {user_query}"}
            ]
            return get_ai_response(messages)

        else:
            return "❌ Please upload PDFs or set a research topic."

    except Exception as e:
        return f"[RAG ERROR] {e}"

# ------------------------------
# PDF Utilities
# ------------------------------
def extract_text_from_pdfs(pdf_files):
    texts, names = [], []
    for pdf_file in pdf_files:
        try:
            doc = fitz.open(pdf_file.name)
            pdf_text = "".join([page.get_text() for page in doc])
            texts.append(pdf_text)
            names.append(os.path.basename(pdf_file.name))
        except Exception as e:
            print(f"Error reading PDF: {e}")
    return texts, names

# ------------------------------
# arXiv Utilities
# ------------------------------
def fetch_research_papers(topic, max_results=5):
    papers = []
    try:
        search = arxiv.Search(query=topic, max_results=max_results, sort_by=arxiv.SortCriterion.Relevance)
        for result in search.results():
            papers.append({
                "title": result.title,
                "authors": [a.name for a in result.authors],
                "abstract": result.summary,
                "pdf_url": result.pdf_url
            })
    except Exception as e:
        print("arXiv fetch error:", e)
    return papers

# ------------------------------
# Chat Integration
# ------------------------------
def combined_chat(idea_topic, pdf_files, user_input):
    global current_research_topic, has_uploaded_pdfs

    if not user_input:
        return "❌ Please enter a question."

    if idea_topic:
        current_research_topic = idea_topic

    if pdf_files and not has_uploaded_pdfs:
        pdf_texts, pdf_names = extract_text_from_pdfs(pdf_files)
        if pdf_texts:
            add_texts_to_rag(pdf_texts, pdf_names, source_type="pdf")
            has_uploaded_pdfs = True

    session_messages.append({"role": "user", "content": user_input})
    answer = query_rag(user_input, research_topic=current_research_topic)
    session_messages.append({"role": "assistant", "content": answer})

    return answer

# ------------------------------
# Enhanced Paper Generator with Quality Processing
# ------------------------------
def generate_individual_section(section_name, topic, context_info, abstracts):
    """Generate a specific section with detailed prompts"""
    
    section_prompts = {
        "Abstract": f"""
        Write a comprehensive abstract (150-250 words) for a research paper on '{topic}'. 
        Include: background, objectives, methodology, key findings, and conclusions.
        Make it concise but informative.
        """,
        
        "Introduction": f"""
        Write a detailed introduction for a research paper on '{topic}'. 
        Include: problem statement, research objectives, significance, and paper structure.
        Provide context and motivation for the research.
        """,
        
        "Literature Review": f"""
        Write a comprehensive literature review for a research paper on '{topic}'. 
        Analyze existing research, identify gaps, and position your work.
        Include relevant studies and theoretical frameworks.
        """,
        
        "Methodology": f"""
        Write a detailed methodology section for a research paper on '{topic}'. 
        Describe research design, data collection methods, analysis techniques, and tools used.
        Be specific about procedures and justify choices.
        """,
        
        "Results and Discussion": f"""
        Write a results and discussion section for a research paper on '{topic}'. 
        Present findings with analysis, compare with existing literature, 
        discuss implications, limitations, and future work.
        """,
        
        "Conclusion": f"""
        Write a conclusion for a research paper on '{topic}'. 
        Summarize key findings, contributions, limitations, and future research directions.
        Be concise but comprehensive.
        """,
        
        "References": f"""
        Generate a reference list for a research paper on '{topic}'. 
        Include 15-20 relevant academic sources (journals, conferences, books).
        Use proper academic citation format.
        """
    }
    
    prompt = section_prompts.get(section_name, f"Write a {section_name} section for a research paper on '{topic}'.")
    
    messages = [
        {"role": "system", "content": f"You are an expert academic writer specializing in {topic}. Write high-quality, original content."},
        {"role": "user", "content": f"{prompt}\n\nContext: {context_info}\n\nRelevant Abstracts: {abstracts}"}
    ]
    
    return get_ai_response(messages)

def parse_paper_sections(paper_text):
    """Parse paper into canonical sections"""
    required = ["Abstract", "Introduction", "Literature Review", "Methodology",
                "Results and Discussion", "Conclusion", "References"]
    sections = {}
    
    # Normalize newlines
    text = re.sub(r'\r\n?', '\n', paper_text).strip()

    # Look for headings that are on their own line
    heading_pattern = re.compile(r'^\s*(Abstract|Introduction|Literature Review|Related Work|Background|Methodology|Methods|Materials and Methods|Results|Results and Discussion|Discussion|Conclusion|Conclusions|References)\s*[:\n\r\-]{0,2}$', re.IGNORECASE | re.MULTILINE)

    matches = list(heading_pattern.finditer(text))
    if matches:
        # Collect sections by heading span
        for i, m in enumerate(matches):
            head = m.group(1).strip()
            start = m.end()
            end = matches[i+1].start() if i+1 < len(matches) else len(text)
            content = text[start:end].strip()
            
            # Normalize headings to canonical names
            if head.lower() in ["related work", "background"]:
                head = "Literature Review"
            if head.lower() in ["methods", "materials and methods"]:
                head = "Methodology"
            if head.lower() == "results":
                head = "Results and Discussion"
            if head.lower() == "conclusions":
                head = "Conclusion"
                
            sections[head] = content
    else:
        # Fallback: split by double newlines and try to map parts
        parts = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        
        for i, part in enumerate(parts):
            if i == 0 and len(part.split()) < 300:
                sections.setdefault("Abstract", part)
            elif i == 0:
                sections.setdefault("Introduction", part)
            elif "method" in part.lower() or "experiment" in part.lower():
                sections.setdefault("Methodology", part)
            elif "result" in part.lower() or "discussion" in part.lower():
                sections.setdefault("Results and Discussion", part)
            elif "conclusion" in part.lower():
                sections.setdefault("Conclusion", part)
            elif "reference" in part.lower():
                sections.setdefault("References", part)
            else:
                if "Introduction" not in sections:
                    sections.setdefault("Introduction", part)
                else:
                    sections.setdefault("Literature Review", (sections.get("Literature Review","") + "\n\n" + part).strip())

    # Ensure canonical order and presence of required keys
    ordered = {}
    for key in ["Abstract", "Introduction", "Literature Review", "Methodology", "Results and Discussion", "Conclusion", "References"]:
        if key in sections and sections[key].strip():
            ordered[key] = sections[key]
        else:
            ordered[key] = ""  # placeholder empty
    return ordered

def generate_paper_with_quality_check(topic, enable_quality_check=True, progress=gr.Progress()):
    """
    Enhanced paper generation with comprehensive quality control:
    1. Generate complete paper structure
    2. Ensure all sections have content
    3. Run quality checks on each section
    4. Humanize and correct as needed
    5. Generate final files
    """
    if not topic:
        return "❌ Please enter a research topic.", None, None, ""

    print(f"🚀 Starting paper generation for: {topic}")
    
    progress(0, desc="Fetching research papers...")
    papers = fetch_research_papers(topic, max_results=5)
    abstracts = " ".join([p["abstract"] for p in papers if p.get("abstract")])
    context_info = f"Research topic: {topic}. Found {len(papers)} relevant papers."

    progress(0.1, desc="Generating initial paper structure...")
    
    # First, try to generate the complete paper
    system_prompt = """You are an expert academic writer. Generate a comprehensive research paper with all required sections. 
    Use clear section headings and ensure each section has substantial content. Be original and academic in tone."""
    
    user_prompt = f"""
    Create a comprehensive research paper on '{topic}' with these sections:
    
    1. Abstract (150-250 words) - Include background, objectives, methodology, findings, conclusions
    2. Introduction - Problem statement, objectives, significance, paper structure
    3. Literature Review - Existing research, gaps, theoretical frameworks
    4. Methodology - Research design, methods, procedures, tools
    5. Results and Discussion - Findings, analysis, implications, limitations
    6. Conclusion - Summary, contributions, future work
    7. References - 15-20 academic sources
    
    Context: {context_info}
    Relevant Abstracts: {abstracts}
    
    Return the complete paper with clear section headings.
    """
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    paper_text = get_ai_response(messages)
    if not paper_text or paper_text.startswith("[AI ERROR]"):
        return f"❌ Could not generate initial paper: {paper_text}", None, None, ""

    progress(0.2, desc="Parsing and validating sections...")
    sections = parse_paper_sections(paper_text)

    # Check which sections need content and generate them individually
    progress(0.3, desc="Ensuring all sections have content...")
    
    for section_name, content in sections.items():
        if not content or len(content.strip()) < 50:
            print(f"📝 Generating missing content for: {section_name}")
            progress(0.3 + 0.1 * (list(sections.keys()).index(section_name) / len(sections)), 
                    desc=f"Generating {section_name}...")
            
            new_content = generate_individual_section(section_name, topic, context_info, abstracts)
            if new_content and not new_content.startswith("[AI ERROR]"):
                sections[section_name] = new_content
                print(f"✅ Generated {section_name}")
            else:
                # Fallback content
                sections[section_name] = f"[PLACEHOLDER] {section_name} content for {topic} - This section requires manual completion."
                print(f"⚠️ Using placeholder for {section_name}")
            
            time.sleep(1)  # Rate limiting

    # Quality processing
    progress(0.5, desc="Starting quality processing...")
    processed_sections = {}
    section_reports = {}
    
    for idx, (section_name, section_content) in enumerate(sections.items()):
        # Skip references from heavy processing
        if section_name.lower() == "references" or not enable_quality_check:
            processed_sections[section_name] = section_content
            section_reports[section_name] = {
                "ai_score": 0, 
                "plagiarism_score": 0, 
                "grammar_issues": 0, 
                "status": "skipped" if section_name.lower()=="references" else "not-checked"
            }
            continue

        progress(0.5 + 0.4 * (idx / max(1, len(sections)-1)), 
                desc=f"Quality processing {section_name}...")
        
        processed_content, report = process_section_quality(section_content, section_name)
        processed_sections[section_name] = processed_content
        section_reports[section_name] = report
        
        time.sleep(0.5)  # Rate limiting

    # Reconstruct the paper in canonical order
    progress(0.9, desc="Assembling final paper...")
    ordered_text = "\n\n".join([
        f"{name}\n{processed_sections.get(name,'')}".strip() 
        for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]
    ])

    # Create files
    progress(0.95, desc="Generating files...")
    os.makedirs("generated_papers", exist_ok=True)
    safe_name = re.sub(r'[^\w\-\. ]', '', topic)[:50]

    # DOCX
    docx_path = f"generated_papers/{safe_name}.docx"
    doc = Document()
    doc.add_heading(topic, level=1)
    
    for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]:
        content = processed_sections.get(name, "")
        if content:
            doc.add_heading(name, level=2)
            # Split content into paragraphs
            paragraphs = re.split(r'\n{2,}', content)
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
    
    doc.save(docx_path)

    # PDF
    pdf_path = f"generated_papers/{safe_name}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50
    
    # Title
    c.setFont("Times-Bold", 16)
    c.drawString(50, y_position, topic)
    y_position -= 30
    
    for name in ["Abstract","Introduction","Literature Review","Methodology","Results and Discussion","Conclusion","References"]:
        content = processed_sections.get(name, "")
        if content:
            # Section heading
            c.setFont("Times-Bold", 12)
            c.drawString(50, y_position, name)
            y_position -= 20
            
            # Section content
            c.setFont("Times-Roman", 10)
            lines = content.split('\n')
            for line in lines:
                if y_position < 50:  # New page
                    c.showPage()
                    y_position = height - 50
                
                # Word wrap
                words = line.split()
                current_line = ""
                for word in words:
                    if c.stringWidth(current_line + word + " ") < width - 100:
                        current_line += word + " "
                    else:
                        if current_line:
                            c.drawString(50, y_position, current_line.strip())
                            y_position -= 12
                        current_line = word + " "
                
                if current_line:
                    c.drawString(50, y_position, current_line.strip())
                    y_position -= 12
                
                y_position -= 5  # Space between paragraphs
    
    c.save()

    # Build quality report
    quality_report_text = ""
    if enable_quality_check:
        quality_report_text = "📊 QUALITY CHECK REPORT\n" + "="*60 + "\n\n"
        for section, report in section_reports.items():
            quality_report_text += f"📄 {section}:\n"
            quality_report_text += f"   AI Score: {report['ai_score']:.1f}%\n"
            quality_report_text += f"   Plagiarism: {report['plagiarism_score']:.1f}%\n"
            quality_report_text += f"   Grammar Issues: {report['grammar_issues']}\n"
            quality_report_text += f"   Status: {report['status']}\n\n"

    progress(1.0, desc="Complete!")
    print(f"✅ Paper generation completed: {topic}")
    
    return ordered_text, docx_path, pdf_path, quality_report_text

# ------------------------------
# Analytics Functions
# ------------------------------
def get_session_statistics():
    total_docs = len(session_docs)
    total_messages = len(session_messages)
    user_messages = len([m for m in session_messages if m["role"]=="user"])
    ai_messages = len([m for m in session_messages if m["role"]=="assistant"])
    total_chars = sum(len(doc["text"]) for doc in session_docs)
    avg_doc_length = total_chars / total_docs if total_docs > 0 else 0
    pdf_count = len([d for d in session_docs if d.get("source") == "pdf"])
    arxiv_count = len([d for d in session_docs if d.get("source") == "arxiv"])

    return {
        "total_documents": total_docs,
        "pdf_documents": pdf_count,
        "arxiv_documents": arxiv_count,
        "total_messages": total_messages,
        "user_messages": user_messages,
        "ai_messages": ai_messages,
        "total_characters": total_chars,
        "avg_doc_length": avg_doc_length
    }

def get_top_keywords(top_n=15):
    all_text = " ".join([doc["text"] for doc in session_docs])
    words = re.findall(r'\b\w+\b', all_text.lower())
    stopwords = set(["the","and","for","with","that","this","from","are","was","were","have","has","using","our","can","which","these","their","been","into","than","more","also","will","such","when","there","other","through","about","some","only","would","between"])
    filtered_words = [w for w in words if w not in stopwords and len(w)>3]
    word_counts = Counter(filtered_words)
    return word_counts.most_common(top_n)

def plot_keyword_bar():
    keywords = get_top_keywords()
    if not keywords:
        fig = go.Figure()
        fig.update_layout(title="Top Keywords", plot_bgcolor="white")
        return fig
    df = pd.DataFrame(keywords, columns=["Keyword","Count"])
    fig = px.bar(df, x="Keyword", y="Count", title="Top Keywords", color="Count")
    return fig

def generate_insights_report_session():
    stats = get_session_statistics()
    keywords = get_top_keywords(10)

    topic_info = f"\n🎯 Topic: {current_research_topic}" if current_research_topic else ""
    source_info = f"\n📄 Source: {'PDFs' if has_uploaded_pdfs else 'arXiv'}"

    report = f"""
📊 Session Insights
{'='*60}
{topic_info}{source_info}

📚 Documents: {stats['total_documents']} ({stats['pdf_documents']} PDFs, {stats['arxiv_documents']} arXiv)
💬 Interactions: {stats['total_messages']}
📝 Characters: {stats['total_characters']:,}

🔑 Keywords: {', '.join([k for k,c in keywords[:10]]) if keywords else 'None'}
"""
    return report

def refresh_all_analytics():
    report = generate_insights_report_session()
    keyword_plot = plot_keyword_bar()
    return report, keyword_plot

# ------------------------------
# Gradio UI
# ------------------------------
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("""
    # 🧠 ResearchMate AI — Advanced Research Assistant with Quality Control
    ### 📚 Upload PDFs OR Set Topic | ✅ AI Detection, Plagiarism Check, Humanization & Grammar Correction
    """)

    with gr.Tabs():
        with gr.Tab("💬 Chat Assistant"):
            with gr.Row():
                clear_btn = gr.Button("🔄 Clear Session", variant="secondary")

            idea_input = gr.Textbox(label="🎯 Research Topic", placeholder="e.g., Machine Learning in Healthcare")
            pdf_upload = gr.Files(label="📄 Upload PDFs (Optional)", file_types=[".pdf"])
            pdf_status = gr.Textbox(label="📋 Status", interactive=False, lines=3)
            chat_input = gr.Textbox(label="❓ Your Question", placeholder="Ask anything...", lines=2)
            output_box = gr.Textbox(label="🤖 AI Answer", interactive=False, lines=12)
            btn = gr.Button("🚀 Send", variant="primary")

            def clear_session():
                global session_messages, session_docs, current_research_topic, has_uploaded_pdfs
                session_messages, session_docs = [], []
                current_research_topic, has_uploaded_pdfs = None, False
                try:
                    collection.delete(where={})
                except:
                    pass
                return "✅ Session cleared!", ""

            clear_btn.click(clear_session, outputs=[pdf_status, output_box])

            def upload_handler(idea_topic, pdf_files):
                global current_research_topic, has_uploaded_pdfs
                if idea_topic:
                    current_research_topic = idea_topic
                if not pdf_files:
                    return f"✅ Topic set: '{current_research_topic}'" if current_research_topic else "⚠️ Set topic or upload PDFs"
                try:
                    pdf_texts, pdf_names = extract_text_from_pdfs(pdf_files)
                    if pdf_texts:
                        add_texts_to_rag(pdf_texts, pdf_names, source_type="pdf")
                        has_uploaded_pdfs = True
                        return f"✅ Processed {len(pdf_texts)} PDF(s)"
                    return "❌ Failed to extract text"
                except Exception as e:
                    return f"❌ Error: {str(e)}"

            pdf_upload.change(upload_handler, inputs=[idea_input, pdf_upload], outputs=pdf_status)
            btn.click(combined_chat, inputs=[idea_input, pdf_upload, chat_input], outputs=output_box)

        with gr.Tab("📝 Paper Generator with Quality Control"):
            gr.Markdown("""
            Generate research papers with automatic quality checking:
            - ✅ AI Content Detection
            - ✅ Plagiarism Checking  
            - ✅ Text Humanization
            - ✅ Grammar Correction
            """)

            topic_box = gr.Textbox(label="Research Topic", placeholder="e.g., Quantum Computing in Cryptography")
            quality_checkbox = gr.Checkbox(label="Enable Quality Processing (AI Detection, Plagiarism, Humanization, Grammar)", value=True)

            paper_output = gr.Textbox(label="Generated Paper", lines=20, interactive=False)
            quality_report_box = gr.Textbox(label="📊 Quality Report", lines=10, interactive=False)

            with gr.Row():
                docx_download = gr.File(label="📥 Download DOCX")
                pdf_download = gr.File(label="📥 Download PDF")

            gen_btn = gr.Button("✍️ Generate Paper", variant="primary")

            gen_btn.click(
                generate_paper_with_quality_check,
                inputs=[topic_box, quality_checkbox],
                outputs=[paper_output, docx_download, pdf_download, quality_report_box]
            )

        with gr.Tab("📊 Analytics"):
            refresh_btn = gr.Button("🔄 Refresh", variant="primary")
            insights_text = gr.Textbox(label="Overview", lines=15, interactive=False)
            keyword_plot_box = gr.Plot(label="Keywords")

            refresh_btn.click(refresh_all_analytics, outputs=[insights_text, keyword_plot_box])
            demo.load(refresh_all_analytics, outputs=[insights_text, keyword_plot_box])

demo.launch(share=True)
